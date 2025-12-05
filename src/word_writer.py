"""
Moduł do generowania dokumentów Word z równaniami.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from .equation_parser import EquationResult


class WordWriter:
    """Klasa do generowania dokumentów Word z równaniami."""
    
    DEFAULT_PRECISION = 4  # Domyślna liczba miejsc po przecinku
    
    def __init__(self, output_path: str | Path, precision: int | None = None):
        """
        Inicjalizuje generator dokumentów Word.
        
        Args:
            output_path: Ścieżka do pliku wyjściowego
            precision: Liczba miejsc po przecinku w wynikach (domyślnie 4)
        """
        self.output_path = Path(output_path)
        self.precision = precision if precision is not None else self.DEFAULT_PRECISION
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self) -> None:
        """Konfiguruje podstawowe ustawienia dokumentu."""
        # Dodaj tytuł dokumentu
        title = self.doc.add_heading("Obliczenia", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def add_equation_section(self, equation_result: EquationResult) -> None:
        """
        Dodaje sekcję z równaniem do dokumentu.
        
        Args:
            equation_result: Wynik przetworzenia równania
        """
        # Nagłówek z nazwą równania
        self.doc.add_heading(equation_result.name, level=1)
        
        # Wzór oryginalny
        p1 = self.doc.add_paragraph()
        p1.add_run("Wzór: ").bold = True
        p1.add_run(equation_result.original_equation)
        
        # Wzór z podstawionymi wartościami
        p2 = self.doc.add_paragraph()
        p2.add_run("Po podstawieniu: ").bold = True
        p2.add_run(equation_result.equation_with_values)
        
        # Wynik
        p3 = self.doc.add_paragraph()
        p3.add_run("Wynik: ").bold = True
        p3.add_run(f"{equation_result.result:.{self.precision}f}")
        
        # Pusty paragraf jako separator
        self.doc.add_paragraph()
    
    def add_variables_table(self, variables: dict) -> None:
        """
        Dodaje tabelę ze zmiennymi do dokumentu.
        
        Args:
            variables: Słownik ze zmiennymi i ich wartościami
        """
        self.doc.add_heading("Dane wejściowe", level=1)
        
        # Utwórz tabelę
        table = self.doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        
        # Nagłówki tabeli
        header_cells = table.rows[0].cells
        header_cells[0].text = "Zmienna"
        header_cells[1].text = "Wartość"
        
        # Dane
        for var_name, value in variables.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(var_name)
            row_cells[1].text = str(value)
        
        self.doc.add_paragraph()
    
    def add_results_section(self, results: list[EquationResult]) -> None:
        """
        Dodaje sekcję z wynikami obliczeń.
        
        Args:
            results: Lista wyników równań
        """
        self.doc.add_heading("Wyniki obliczeń", level=1)
        
        for result in results:
            self.add_equation_section(result)
    
    def save(self) -> Path:
        """
        Zapisuje dokument do pliku.
        
        Returns:
            Ścieżka do zapisanego pliku
        """
        self.doc.save(self.output_path)
        return self.output_path
