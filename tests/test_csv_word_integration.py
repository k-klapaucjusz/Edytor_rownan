"""
Testy integracyjne dla przetwarzania plików CSV i generowania dokumentów Word.
"""

import os
import tempfile
from pathlib import Path

import pytest
from docx import Document

from src.excel_reader import CSVReader
from src.main import process_csv_equations


class TestCSVReader:
    """Testy dla klasy CSVReader."""
    
    def test_read_variables_from_csv(self):
        """Test wczytywania zmiennych z pliku CSV."""
        csv_path = Path(__file__).parent.parent / "examples" / "dane_prad_3fazowy.csv"
        reader = CSVReader(csv_path)
        variables = reader.read_variables()
        
        assert "P" in variables
        assert "U" in variables
        assert "cos_phi" in variables
        assert variables["P"] == 15000
        assert variables["U"] == 400
        assert variables["cos_phi"] == 0.85
    
    def test_read_variables_with_metadata(self):
        """Test wczytywania zmiennych z metadanymi."""
        csv_path = Path(__file__).parent.parent / "examples" / "dane_prad_3fazowy.csv"
        reader = CSVReader(csv_path)
        data = reader.read_variables_with_metadata()
        
        assert len(data) == 3
        # Sprawdź czy zawiera jednostki
        assert any(item.get("Jednostka") == "W" for item in data)
    
    def test_file_not_found(self):
        """Test obsługi błędu gdy plik nie istnieje."""
        with pytest.raises(FileNotFoundError):
            CSVReader("nieistniejacy_plik.csv")
    
    def test_insufficient_columns(self):
        """Test obsługi błędu gdy plik CSV ma za mało kolumn."""
        # Utwórz tymczasowy plik CSV z jedną kolumną
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as tmp:
            tmp.write("Wartość\n")
            tmp.write("100\n")
            tmp_path = tmp.name
        
        try:
            reader = CSVReader(tmp_path)
            with pytest.raises(ValueError, match="co najmniej 2 kolumny"):
                reader.read_variables()
        finally:
            os.remove(tmp_path)


class TestProcessCsvEquations:
    """Testy dla funkcji process_csv_equations."""
    
    def test_generate_word_from_csv(self):
        """Test generowania dokumentu Word z danych CSV."""
        csv_path = Path(__file__).parent.parent / "examples" / "dane_prad_3fazowy.csv"
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = tmp.name
        
        try:
            equations = [
                {"name": "Prąd fazowy", "formula": "P / (sqrt(3) * U * cos_phi)"}
            ]
            
            result_path = process_csv_equations(
                csv_path,
                output_path,
                equations,
                title="Obliczenia prądu 3-fazowego"
            )
            
            assert Path(result_path).exists()
            
            # Sprawdź zawartość dokumentu
            doc = Document(result_path)
            
            # Sprawdź tytuł
            assert doc.paragraphs[0].text == "Obliczenia prądu 3-fazowego"
            
            # Sprawdź czy są tabele (dane wejściowe)
            assert len(doc.tables) >= 1
            
            # Sprawdź czy wynik jest poprawny (~25.47 A)
            # Używamy zakresu 25.4 - 25.5 dla stabilności testu
            text_content = "\n".join(p.text for p in doc.paragraphs)
            import re
            result_match = re.search(r"Wynik:\s*([\d.]+)", text_content)
            assert result_match is not None, "Nie znaleziono wyniku w dokumencie"
            result_value = float(result_match.group(1))
            assert 25.4 < result_value < 25.5, f"Wynik {result_value} poza oczekiwanym zakresem 25.4-25.5"
            
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)
    
    def test_multiple_equations(self):
        """Test przetwarzania wielu równań."""
        csv_path = Path(__file__).parent.parent / "examples" / "dane_prad_3fazowy.csv"
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = tmp.name
        
        try:
            equations = [
                {"name": "Prąd fazowy", "formula": "P / (sqrt(3) * U * cos_phi)"},
                {"name": "Moc pozorna", "formula": "P / cos_phi"},
            ]
            
            result_path = process_csv_equations(
                csv_path,
                output_path,
                equations,
            )
            
            doc = Document(result_path)
            text_content = "\n".join(p.text for p in doc.paragraphs)
            
            # Sprawdź czy oba równania są obecne
            assert "Prąd fazowy" in text_content
            assert "Moc pozorna" in text_content
            
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)
